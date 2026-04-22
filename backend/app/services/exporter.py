"""
Export comparison results as PDF or DOCX redline.
"""
import io
from typing import List


def export_pdf(original_spans: List[dict], revised_spans: List[dict], comparison_id: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontSize=16, spaceAfter=12)
    normal = styles["Normal"]

    COLOR_DELETE = colors.HexColor("#ffcccc")
    COLOR_INSERT = colors.HexColor("#ccffcc")
    COLOR_MODIFY = colors.HexColor("#fff3cc")

    def build_para(spans, side):
        parts = []
        for span in spans:
            t = span["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not t:
                continue
            stype = span["type"]
            if stype == "delete" and side == "original":
                parts.append(f'<font backColor="#ffcccc"><strike>{t}</strike></font>')
            elif stype == "insert" and side == "revised":
                parts.append(f'<font backColor="#ccffcc">{t}</font>')
            elif stype == "replace":
                if side == "original":
                    parts.append(f'<font backColor="#fff3cc"><strike>{t}</strike></font>')
                else:
                    parts.append(f'<font backColor="#fff3cc">{t}</font>')
            else:
                parts.append(t)
        return Paragraph("".join(parts), normal)

    story = [
        Paragraph(f"DocCompare Pro — Comparison Report", title_style),
        Paragraph(f"Comparison ID: {comparison_id}", normal),
        Spacer(1, 0.5*cm),
    ]

    orig_para = build_para(original_spans, "original")
    rev_para  = build_para(revised_spans,  "revised")

    tbl = Table([[
        Paragraph("<b>Original</b>", styles["Heading3"]),
        Paragraph("<b>Revised</b>",  styles["Heading3"]),
    ], [orig_para, rev_para]], colWidths=["50%", "50%"])

    tbl.setStyle(TableStyle([
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND",  (0, 0), (-1,  0), colors.HexColor("#1F497D")),
        ("TEXTCOLOR",   (0, 0), (-1,  0), colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f8f8")]),
    ]))
    story.append(tbl)

    doc.build(story)
    return buf.getvalue()


def export_docx(original_spans: List[dict], revised_spans: List[dict]) -> bytes:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()
    document.add_heading("DocCompare Pro — Redline Document", level=1)

    def add_highlighted_paragraph(doc, spans, side, heading_text):
        doc.add_heading(heading_text, level=2)
        para = doc.add_paragraph()
        for span in spans:
            text = span["text"]
            if not text:
                continue
            stype = span["type"]
            run = para.add_run(text)
            if stype == "delete" and side == "original":
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.strike = True
                _set_highlight(run, "red")
            elif stype == "insert" and side == "revised":
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                _set_highlight(run, "green")
            elif stype == "replace":
                run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
                _set_highlight(run, "yellow")
                if side == "original":
                    run.font.strike = True

    def _set_highlight(run, color_name):
        rPr = run._r.get_or_add_rPr()
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), color_name)
        rPr.append(highlight)

    add_highlighted_paragraph(document, original_spans, "original", "Original Document")
    document.add_page_break()
    add_highlighted_paragraph(document, revised_spans,  "revised",  "Revised Document")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
